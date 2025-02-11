#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time    : 2022/12/30
# @Author  : browser
# @Software: PyCharm
# @File    : cuttlefish.py


import argparse
import re
import time
import os
import json
import openai
import backoff
from DecryptLogin import login
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

cid_name: dict = {0: '学前教育', 1: '基础教育', 2: '高校与高等教育', 3: '语言/资格考试', 4: '法律', 5: '建筑', 6: '互联网', 7: '行业资料', 8: '政务民生',
                  9: '商品说明书', 10: '实用模板', 11: '生活娱乐', 99: '推荐'}

# Set your API key
openai.api_key = "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"


def logging(msg, tip='INFO'):
    """
    日志格式化
    :param msg: 格式化信息
    :param tip: 日志等级
    :return:
    """
    print(f'[{time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())} {tip}]: {msg}')


def get_login():
    """
    百度扫码登入，返回session
    :return: session
    """
    lg = login.Login()
    infos_return, session = lg.baidu('', '', 'scanqr')
    return session


def parseArgs():
    """
    命令行参数解析
    :return: args
    """
    parser = argparse.ArgumentParser(description='墨斗鱼文章生成')
    parser.add_argument('--refresh', dest='refresh', help='强制刷新任务列表', action='store_true')
    parser.add_argument('--n', dest='cid_num', type=int, default=6,
                        help='default=6, 分类编号: 0: 学前教育, 1: 基础教育, 2: 高校与高等教育, 3: 语言/资格考试, 4: 法律, 5: 建筑, 6: 互联网, 7: 行业资料, 8: 政务民生, 9: 商品说明书, 10: 实用模板, 11: 生活娱乐')
    parser.add_argument('--all', dest='all', help='分类全选', action='store_true')
    args = parser.parse_args()
    return args


def traditional_chinese_check(chinese_char):
    """
    生僻字和繁体字过滤
    使用 `in` 操作符判断中文字符是否存在于文件中
    :return: true or false
    """
    with open('rc.txt', 'r', encoding='UTF-8') as f:
        content = f.read()
    flag = False
    n = 0
    while n < len(chinese_char):
        for i in chinese_char:
            if i in content:
                flag = True
                break
            else:
                flag = False
                n += 1
        if flag:
            break
    if flag:
        return True
    else:
        return False


class Gater_wool:
    """
    该死，就是想薅一把羊毛，行动起来
    """

    def __init__(self):
        self.sum = 200  # 每天api接口调用上线
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/108.0.0.0 Safari/537.36',
            'Content-Type': 'application/x-www-form-urlencoded',
            'Referer': 'https://cuttlefish.baidu.com/shopmis',
            'Accept': '*/*'
        }

    @property
    def get_task(self):
        """
        获取墨斗鱼任务列表
        :return: list
        """
        args = parseArgs()
        if args.all:
            cid_list: list = [0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11]
        elif args.cid_num or int(args.cid_num == 0):
            cid_list: list = [args.cid_num]
        else:
            cid_list: list = [99]
        task_name: list = []
        total_name: dict = {}
        session = get_login()
        for cid in cid_list:
            pn: int = 0
            rn: int = 20
            f_url = f'https://cuttlefish.baidu.com/user/interface/getquerypacklist?cid={cid}&pn={pn}&rn={rn}&word=&tab=1'
            response = session.get(f_url, headers=self.headers)
            response_json = response.json()
            if response_json['status']['code'] != 0:
                logging('%s 第 %s 页，获取任务名失败, 原因: %s' % (cid_name[cid], pn + 1, response_json.get('status')))
            else:
                total = response_json['data']['total']
                pn = int(total / rn)
                t_pn: int = 0
                while t_pn < pn:
                    t_url = f'https://cuttlefish.baidu.com/user/interface/getquerypacklist?cid={cid}&pn={t_pn}&rn={rn}&word=&tab=1'
                    t_response = session.get(t_url, headers=self.headers)
                    t_response_json = t_response.json()
                    if t_response_json['status']['code'] != 0:
                        logging('%s 第 %s 页，获取任务名失败, 原因: %s' % (cid_name[cid], t_pn, response_json.get('status')))
                    else:
                        # 从响应内容中获取任务名
                        for i in t_response_json['data']['queryList']:
                            if int(i['status']) == 1:  # 过滤已完成的标题
                                t_name = i['queryName']
                                if len(i['queryName']) < 5:  # 标题长度不得小于5个字
                                    logging('%s 标题长度小于5个字符---不符' % t_name)
                                elif re.search(r"[a-zA-Z0-9\*\"/:?\\|<>x×÷+“：《》～,，\-.()。\ ·_#$!%&@【】^{}~]+",
                                               t_name):  # 过滤特殊字符
                                    logging('%s 包含特殊字符---不符' % t_name)
                                elif traditional_chinese_check(t_name):
                                    logging('%s 包含繁体字和生僻字---不符' % t_name)
                                else:
                                    task_name.append(t_name)

                        # 判断是否成功
                        if t_response_json['status']['code'] == 0:
                            logging('%s 第 %s 页，获取任务名成功 %s' % (cid_name[cid], t_pn + 1, response_json.get('status')))
                        else:
                            logging(
                                '%s 第 %s 页，获取任务名失败, 原因: %s' % (cid_name[cid], t_pn + 1, response_json.get('status')))
                    t_pn += 1
                    time.sleep(1.5)
            total_name[cid_name[cid]] = task_name
        return total_name

    def write_to_file(self):
        """
        任务名信息存入文件
        :return: file
        """
        args = parseArgs()
        if args.refresh:
            task_info = self.get_task
            f = open('../task/task_name.json', 'w', encoding='utf-8')
            json.dump(task_info, f, ensure_ascii=False, indent=4)
            logging('任务名刷新并写入成功!!!')
            f.close()
            return task_info
        elif os.path.exists('../task/task_name.json'):
            f = open('../task/task_name.json', 'r', encoding='utf-8')
            task_info = json.load(f)
            logging('任务名成功读取!!!')
            f.close()
            return task_info
        else:
            task_info = self.get_task
            f = open('../task/task_name.json', 'w', encoding='utf-8')
            json.dump(task_info, f, ensure_ascii=False, indent=4)
            logging('任务名首次写入成功!!!')
            f.close()
            return task_info

    def write_to_docx(self, classification, title, content):
        """
        生成docx的word文档
        :return: .docx
        """
        # 打开文档
        document = Document()

        # 标题
        head = document.add_heading()
        head.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        head.paragraph_format.space_before = Pt(36)
        run = head.add_run(title)
        run.font.size = Pt(36)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 标题颜色
        document.styles['Normal'].font.name = u'Times New Roman'  # 设置西文字体
        document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置中文字体使用字体2->宋体
        # 添加分页
        # document.add_page_break()

        # 二级标题
        head2 = document.add_heading(level=2)
        head2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = head2.add_run(u'第一章')
        run.font.name = u'宋体'
        run.font.size = Pt(21)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        # 段落后行距
        # head2.paragraph_format.space_after = Pt(10)

        # 段落
        p = document.add_paragraph()
        run = p.add_run(content)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(14)

        # 附录
        head3 = document.add_heading(level=2)
        head2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = head3.add_run(u'附录:')
        run.font.name = u'宋体'
        run.font.size = Pt(21)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        # 附录内容
        p = document.add_paragraph()
        run = p.add_run(u'无')
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(14)

        # 单倍行距
        p.paragraph_format.line_spacing = Pt(30)
        # 段落后行距
        p.paragraph_format.space_after = Pt(30)
        # 段落分页
        p.keep_with_next = True
        # 保存
        path = f'../file/{classification}/{time.strftime("%Y-%m-%d")}/'
        filename = f'{title}.docx'
        if not os.path.exists(path):
            os.makedirs(path)
        document.save(path + filename)

    @backoff.on_exception(backoff.expo, openai.error.RateLimitError)
    def completions_with_backoff(self):
        count_sum = self.sum
        num = 0
        total_task_name = self.write_to_file()
        for key, value in total_task_name.items():
            classification = str(key)  # 分类
            for name in value:
                if len(name) < 5:
                    logging('%s 标题长度小于5个字符---不符' % name)
                elif re.search(r"[a-zA-Z0-9\*\"/:?\\|<>x×÷+“：《》～,，\-.()。\ ·_#$!%&@【】^{}~]+", name):
                    logging('%s 包含特殊字符---不符' % name)
                elif traditional_chinese_check(name):
                    logging('%s 包含繁体字和生僻字---不符' % name)
                else:
                    prompt = f'以 {name} 为题，写一篇1000字的文章'
                    logging('第 %s 篇，【%s】 文章生成任务---开始' % (num, prompt))
                    try:
                        # Use the ChatGPT model to generate a response to a prompt
                        response = openai.Completion.create(
                            model="text-davinci-003",
                            prompt=prompt,
                            max_tokens=4000,
                            temperature=0.7,
                            top_p=1,
                            frequency_penalty=0,
                            presence_penalty=0
                        )
                        content = response.choices[0].text
                        self.write_to_docx(classification=classification, title=name, content=content)
                        logging('第 %s 篇，【%s】docx文档---已生成' % (num, name))
                        time.sleep(20)
                    except Exception as e:
                        logging('openai接口异常: %s' % e)
                    # 每天生成篇数限制
                    if num >= count_sum:
                        break
                    num += 1

    # def read_file(self):
    #     """
    #     测试使用
    #     """
    #     total_task_name = self.write_to_file()
    #     a: list = []
    #     b: dict = {}
    #     for key, value in total_task_name.items():
    #         for name in value:
    #             if len(name) < 5:
    #                 logging('%s 标题长度小于5个字符---不符' % name)
    #             elif re.search(r"[a-zA-Z0-9\*\"/:?\\|<>x×÷+“：《》～,，\-.()。\ ·_#$!%&@【】^{}~]+", name):
    #                 logging('%s 包含特殊字符---不符' % name)
    #             elif traditional_chinese_check(name):
    #                 logging('%s 包含繁体字和生僻字---不符' % name)
    #             else:
    #                 a.append(name)
    #         b[key] = a
    #     task_info = b
    #     f = open('../task/b.json', 'w', encoding='utf-8')
    #     json.dump(task_info, f, ensure_ascii=False, indent=4)
    #     logging('过滤成功!!!')
    #     f.close()


if __name__ == '__main__':
    Gater_wool = Gater_wool()
    Gater_wool.completions_with_backoff()
    # Gater_wool.write_to_file()
    # Gater_wool.read_file()
