#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time    : 2023/1/4 16:05
# @Author  : Browser
# @File    : demo_docx.py
# @Software: PyCharm

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

header = '农夫家里出千金'
content = '''
"农夫家里出千金"这个说法指的是一个农夫家里的女儿出嫁，并且嫁得非常好，能够带来巨大的经济收益。在过去，农村家庭通常比较穷，农民也不是很有钱，所以女儿的嫁妆会是家庭的一大支出。如果女儿能够嫁给一个有钱的男人，那么这对家庭来说就是一笔巨大的收益，就像是赚了一千金一样。

现在，这个说法仍然被使用，但通常不再指代女儿嫁妆的意思。它可以用来描述一个人或一个家庭取得巨大的经济成功，就像是从农村走出来，获得了巨大的财富。
'''


def write_to_docx():
    document = Document()
    document.add_heading(header, 1)
    # 段落
    p = document.add_paragraph(content)
    # 段落左对齐
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    # 段落首行缩进2个字符
    # p.first_line_indent = Inches(2)
    # 段落间距
    p.space_before = Pt(0)
    p.space_after = Pt(1)
    # 段落行距
    p.line_spacing_rule = WD_LINE_SPACING.MULTIPLE  # 多倍行距
    p.line_spacing = 1.5  # 1.5倍行间距
    # 段落分页
    p.keep_with_next = True
    # 新增分页符
    # document.add_page_break()
    # 保存
    path_to_file = f'../file//{header}.docx'
    document.save(path_to_file)


# def test():
#     from docx import Document
#     from docx.shared import Inches
#
#     document = Document()
#
#     document.add_heading('Document Title', 0)
#
#     p = document.add_paragraph('A plain paragraph having some ')
#     p.add_run('bold').bold = True
#     p.add_run(' and some ')
#     p.add_run('italic.').italic = True
#
#     document.add_heading('Heading, level 1', level=1)
#     document.add_paragraph('Intense quote', style='Intense Quote')
#
#     document.add_paragraph(
#         'first item in unordered list', style='List Bullet'
#     )
#     document.add_paragraph(
#         'first item in ordered list', style='List Number'
#     )
#
#     document.add_picture('monty-truth.png', width=Inches(1.25))
#
#     records = (
#         (3, '101', 'Spam'),
#         (7, '422', 'Eggs'),
#         (4, '631', 'Spam, spam, eggs, and spam')
#     )
#
#     table = document.add_table(rows=1, cols=3)
#     hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = 'Qty'
#     hdr_cells[1].text = 'Id'
#     hdr_cells[2].text = 'Desc'
#     for qty, id, desc in records:
#         row_cells = table.add_row().cells
#         row_cells[0].text = str(qty)
#         row_cells[1].text = id
#         row_cells[2].text = desc
#
#     document.add_page_break()
#
#     document.save('demo.docx')


if __name__ == '__main__':
    # test()
    write_to_docx()